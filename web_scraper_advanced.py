"""
Web Scraper Avancé - Version améliorée avec gestion des cas bloquants
Auteur: Claude AI
Version: 2.0
Description: Scraper web robuste avec contournement des protections et export MD/Word
"""

import os
import sys
import time
import base64
import logging
import random
from datetime import datetime
from urllib.parse import urlparse, urljoin
from typing import List, Tuple, Dict, Optional
import json

# Selenium et gestion navigateur
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (
    TimeoutException,
    WebDriverException,
    NoSuchElementException,
    StaleElementReferenceException
)
from webdriver_manager.chrome import ChromeDriverManager

# BeautifulSoup pour parsing HTML
from bs4 import BeautifulSoup

# Requests pour téléchargement images
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# Export Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Progress bar
from tqdm import tqdm


# ==================== CONFIGURATION ====================

class ScraperConfig:
    """Configuration centralisée du scraper"""

    # Liste de User-Agents pour rotation
    USER_AGENTS = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:121.0) Gecko/20100101 Firefox/121.0',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.1 Safari/605.1.15'
    ]

    # Timeouts
    PAGE_LOAD_TIMEOUT = 30
    IMPLICIT_WAIT = 10
    EXPLICIT_WAIT = 15
    SCROLL_PAUSE_TIME = 2

    # Retry configuration
    MAX_RETRIES = 3
    RETRY_DELAY = 2  # Secondes
    BACKOFF_FACTOR = 2

    # Dossiers de sortie
    OUTPUT_DIR = "scraped_content"
    IMAGES_DIR = "images"

    # Options de scraping
    HEADLESS = True
    DISABLE_IMAGES = False  # Mettre True pour performance
    ENABLE_JAVASCRIPT = True
    SCROLL_TO_BOTTOM = True  # Pour lazy loading

    # Eléments à supprimer
    UNWANTED_TAGS = ["script", "style", "nav", "footer", "header", "aside", "noscript", "iframe", "meta", "link"]
    UNWANTED_CLASSES = ["advertisement", "ad", "popup", "modal", "cookie-banner", "social-share"]


# ==================== LOGGING ====================

def setup_logging(log_file: str = "scraper.log") -> logging.Logger:
    """Configure le système de logging"""
    logger = logging.getLogger("WebScraper")
    logger.setLevel(logging.INFO)

    # Format
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )

    # Handler console
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(formatter)

    # Handler fichier
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(formatter)

    logger.addHandler(console_handler)
    logger.addHandler(file_handler)

    return logger


# ==================== DRIVER SELENIUM ====================

class AdvancedWebDriver:
    """Gestionnaire de WebDriver Selenium avec options avancées"""

    def __init__(self, config: ScraperConfig, logger: logging.Logger):
        self.config = config
        self.logger = logger
        self.driver = None

    def create_driver(self, proxy: Optional[str] = None, use_stealth: bool = True) -> webdriver.Chrome:
        """Crée un driver Chrome avec options anti-détection"""

        options = Options()

        # Mode headless
        if self.config.HEADLESS:
            options.add_argument('--headless=new')

        # Arguments anti-détection
        options.add_argument('--disable-blink-features=AutomationControlled')
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)

        # User-Agent aléatoire
        user_agent = random.choice(self.config.USER_AGENTS)
        options.add_argument(f'user-agent={user_agent}')
        self.logger.info(f"User-Agent: {user_agent}")

        # Arguments de performance et stabilité
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--disable-gpu')
        options.add_argument('--disable-software-rasterizer')
        options.add_argument('--disable-extensions')
        options.add_argument('--dns-prefetch-disable')

        # Gestion des images (désactiver pour performance)
        if self.config.DISABLE_IMAGES:
            prefs = {
                "profile.managed_default_content_settings.images": 2,
                "profile.default_content_setting_values.notifications": 2
            }
            options.add_experimental_option("prefs", prefs)

        # Proxy
        if proxy:
            options.add_argument(f'--proxy-server={proxy}')
            self.logger.info(f"Utilisation du proxy: {proxy}")

        # SSL
        options.add_argument('--ignore-certificate-errors')
        options.add_argument('--ignore-ssl-errors')

        # Taille fenêtre
        options.add_argument('--window-size=1920,1080')
        options.add_argument('--start-maximized')

        try:
            # Créer le driver
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=options)

            # Paramètres anti-détection supplémentaires
            if use_stealth:
                driver.execute_cdp_cmd('Network.setUserAgentOverride', {
                    "userAgent": user_agent
                })
                driver.execute_script(
                    "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
                )

            # Timeouts
            driver.set_page_load_timeout(self.config.PAGE_LOAD_TIMEOUT)
            driver.implicitly_wait(self.config.IMPLICIT_WAIT)

            self.driver = driver
            self.logger.info("Driver Chrome créé avec succès")
            return driver

        except Exception as e:
            self.logger.error(f"Erreur lors de la création du driver: {e}")
            raise

    def safe_get(self, url: str, max_retries: int = None) -> bool:
        """Charge une page avec retry automatique"""
        if not self.driver:
            raise ValueError("Driver non initialisé. Appelez create_driver() d'abord.")

        max_retries = max_retries or self.config.MAX_RETRIES
        delay = self.config.RETRY_DELAY

        for attempt in range(1, max_retries + 1):
            try:
                self.logger.info(f"Tentative {attempt}/{max_retries} : Chargement de {url}")
                self.driver.get(url)

                # Attendre que la page soit chargée
                WebDriverWait(self.driver, self.config.EXPLICIT_WAIT).until(
                    lambda d: d.execute_script('return document.readyState') == 'complete'
                )

                self.logger.info(f"✓ Page chargée avec succès")
                return True

            except TimeoutException:
                self.logger.warning(f"Timeout lors du chargement (tentative {attempt})")
                if attempt < max_retries:
                    time.sleep(delay)
                    delay *= self.config.BACKOFF_FACTOR

            except WebDriverException as e:
                self.logger.error(f"Erreur WebDriver: {e}")
                if attempt < max_retries:
                    time.sleep(delay)
                    delay *= self.config.BACKOFF_FACTOR
                    # Recréer le driver si nécessaire
                    if "session deleted" in str(e).lower():
                        self.logger.info("Recréation du driver...")
                        self.close()
                        self.create_driver()
                else:
                    raise

        self.logger.error(f"Échec du chargement après {max_retries} tentatives")
        return False

    def handle_popups_and_cookies(self):
        """Ferme les pop-ups et bannières de cookies communes"""
        # Liste de sélecteurs CSS courants pour fermer les pop-ups
        close_selectors = [
            # Boutons de fermeture génériques
            'button[class*="close"]', 'button[class*="dismiss"]',
            'button[aria-label*="close"]', 'button[aria-label*="Close"]',
            '[class*="cookie"] button[class*="accept"]',
            '[class*="cookie"] button[class*="agree"]',
            '[id*="cookie"] button', '[class*="gdpr"] button',
            'button:contains("Accept")', 'button:contains("Accepter")',
            'a[class*="close"]', '.modal-close', '.popup-close'
        ]

        for selector in close_selectors:
            try:
                elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                for element in elements:
                    if element.is_displayed():
                        element.click()
                        self.logger.info(f"Pop-up/Cookie banner fermé: {selector}")
                        time.sleep(0.5)
            except Exception:
                continue

    def scroll_to_bottom(self, pause_time: float = None):
        """Scroll progressif vers le bas (pour lazy loading)"""
        pause_time = pause_time or self.config.SCROLL_PAUSE_TIME

        self.logger.info("Scroll vers le bas pour charger le contenu dynamique...")

        last_height = self.driver.execute_script("return document.body.scrollHeight")
        scroll_attempts = 0
        max_attempts = 10

        while scroll_attempts < max_attempts:
            # Scroll vers le bas
            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(pause_time)

            # Calculer la nouvelle hauteur
            new_height = self.driver.execute_script("return document.body.scrollHeight")

            if new_height == last_height:
                # Plus de contenu à charger
                break

            last_height = new_height
            scroll_attempts += 1

        self.logger.info(f"Scroll terminé après {scroll_attempts} tentatives")

        # Retour en haut
        self.driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(0.5)

    def detect_blocking(self) -> Tuple[bool, str]:
        """Détecte si la page est bloquée (Cloudflare, Captcha, etc.)"""
        page_source = self.driver.page_source.lower()
        title = self.driver.title.lower()

        # Détection Cloudflare
        if any(keyword in page_source for keyword in ['cloudflare', 'cf-browser-verification']):
            return True, "Cloudflare"

        # Détection Captcha
        if any(keyword in page_source for keyword in ['captcha', 'recaptcha', 'hcaptcha']):
            return True, "Captcha"

        # Détection blocage général
        if any(keyword in title for keyword in ['access denied', 'blocked', '403', '429']):
            return True, "Access Denied / Rate Limited"

        # Détection page d'erreur
        if any(keyword in title for keyword in ['error', '404', '500', '502', '503']):
            return True, f"Error Page ({title})"

        return False, ""

    def close(self):
        """Ferme proprement le driver"""
        if self.driver:
            try:
                self.driver.quit()
                self.logger.info("Driver fermé")
            except Exception as e:
                self.logger.error(f"Erreur lors de la fermeture du driver: {e}")


# ==================== EXTRACTION DE CONTENU ====================

class ContentExtractor:
    """Extracteur de contenu avancé"""

    def __init__(self, config: ScraperConfig, logger: logging.Logger):
        self.config = config
        self.logger = logger

    def extract_all(self, html: str, url: str) -> Dict:
        """Extrait tout le contenu de la page"""
        soup = BeautifulSoup(html, 'html.parser')

        self.logger.info("Extraction du contenu...")

        # Nettoyage
        self._clean_soup(soup)

        # Extraction
        title = self._extract_title(soup, url)
        metadata = self._extract_metadata(soup)
        text = self._extract_text(soup)
        tables = self._extract_tables(soup)
        images = self._extract_images(soup, url)
        links = self._extract_links(soup, url)

        return {
            'url': url,
            'title': title,
            'metadata': metadata,
            'text': text,
            'tables': tables,
            'images': images,
            'links': links,
            'extraction_time': datetime.now().isoformat()
        }

    def _clean_soup(self, soup: BeautifulSoup):
        """Nettoie le HTML en supprimant les éléments indésirables"""
        # Supprimer les tags indésirables
        for tag in self.config.UNWANTED_TAGS:
            for element in soup.find_all(tag):
                element.decompose()

        # Supprimer les éléments avec classes indésirables
        for class_name in self.config.UNWANTED_CLASSES:
            for element in soup.find_all(class_=lambda x: x and class_name in x.lower()):
                element.decompose()

    def _extract_title(self, soup: BeautifulSoup, url: str) -> str:
        """Extrait le titre de la page"""
        # Essayer différentes sources
        title = None

        # <title>
        if soup.title and soup.title.string:
            title = soup.title.string.strip()

        # <h1>
        if not title:
            h1 = soup.find('h1')
            if h1:
                title = h1.get_text(strip=True)

        # Open Graph
        if not title:
            og_title = soup.find('meta', property='og:title')
            if og_title and og_title.get('content'):
                title = og_title['content'].strip()

        # Fallback: nom de domaine
        if not title:
            title = urlparse(url).netloc

        return title

    def _extract_metadata(self, soup: BeautifulSoup) -> Dict:
        """Extrait les métadonnées de la page"""
        metadata = {}

        # Description
        description = soup.find('meta', attrs={'name': 'description'})
        if description and description.get('content'):
            metadata['description'] = description['content'].strip()

        # Keywords
        keywords = soup.find('meta', attrs={'name': 'keywords'})
        if keywords and keywords.get('content'):
            metadata['keywords'] = keywords['content'].strip()

        # Author
        author = soup.find('meta', attrs={'name': 'author'})
        if author and author.get('content'):
            metadata['author'] = author['content'].strip()

        # Open Graph
        og_tags = soup.find_all('meta', property=lambda x: x and x.startswith('og:'))
        for tag in og_tags:
            prop = tag.get('property', '').replace('og:', '')
            content = tag.get('content', '')
            if prop and content:
                metadata[f'og_{prop}'] = content.strip()

        return metadata

    def _extract_text(self, soup: BeautifulSoup) -> str:
        """Extrait le texte principal de la page"""
        # Chercher le contenu principal
        main_content = None

        # Essayer les conteneurs courants
        for selector in ['main', 'article', '[role="main"]', '#content', '.content', '#main', '.main']:
            main_content = soup.select_one(selector)
            if main_content:
                break

        # Fallback: body
        if not main_content:
            main_content = soup.find('body')

        if not main_content:
            main_content = soup

        # Extraire le texte
        text = main_content.get_text(separator='\n', strip=True)

        # Nettoyer les lignes vides multiples
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)

        return text

    def _extract_tables(self, soup: BeautifulSoup) -> List[Dict]:
        """Extrait les tableaux"""
        tables = []

        for i, table in enumerate(soup.find_all('table'), 1):
            # Extraire le titre/caption
            caption = table.find('caption')
            table_title = caption.get_text(strip=True) if caption else f"Tableau {i}"

            # Extraire les données
            rows = []
            for row in table.find_all('tr'):
                cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                if cells:  # Ignorer les lignes vides
                    rows.append(cells)

            if rows:
                tables.append({
                    'title': table_title,
                    'rows': rows
                })

        self.logger.info(f"✓ {len(tables)} tableau(x) extrait(s)")
        return tables

    def _extract_images(self, soup: BeautifulSoup, base_url: str) -> List[Dict]:
        """Extrait les informations des images"""
        images = []

        for i, img in enumerate(soup.find_all('img'), 1):
            src = img.get('src') or img.get('data-src') or img.get('data-lazy-src')

            if not src:
                continue

            # URL absolue
            if src.startswith('//'):
                src = 'https:' + src
            elif not src.startswith(('http', 'data:')):
                src = urljoin(base_url, src)

            # Alt text
            alt = img.get('alt', f"Image {i}")

            # Dimensions
            width = img.get('width')
            height = img.get('height')

            images.append({
                'url': src,
                'alt': alt,
                'width': width,
                'height': height,
                'is_data_uri': src.startswith('data:')
            })

        self.logger.info(f"✓ {len(images)} image(s) trouvée(s)")
        return images

    def _extract_links(self, soup: BeautifulSoup, base_url: str) -> List[Dict]:
        """Extrait les liens"""
        links = []
        seen_urls = set()

        for a in soup.find_all('a', href=True):
            href = a['href']
            text = a.get_text(strip=True)

            # URL absolue
            if href.startswith('//'):
                href = 'https:' + href
            elif not href.startswith(('http', 'mailto:', 'tel:', '#')):
                href = urljoin(base_url, href)

            # Éviter les doublons
            if href not in seen_urls and not href.startswith('#'):
                seen_urls.add(href)
                links.append({
                    'url': href,
                    'text': text or href
                })

        self.logger.info(f"✓ {len(links)} lien(s) extrait(s)")
        return links


# ==================== TÉLÉCHARGEMENT D'IMAGES ====================

class ImageDownloader:
    """Gestionnaire de téléchargement d'images"""

    def __init__(self, output_dir: str, logger: logging.Logger):
        self.output_dir = output_dir
        self.logger = logger
        self.session = self._create_session()

        # Créer le dossier si nécessaire
        os.makedirs(output_dir, exist_ok=True)

    def _create_session(self) -> requests.Session:
        """Crée une session requests avec retry"""
        session = requests.Session()

        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504]
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        session.mount("http://", adapter)
        session.mount("https://", adapter)

        session.headers.update({
            'User-Agent': random.choice(ScraperConfig.USER_AGENTS)
        })

        return session

    def download_images(self, images: List[Dict], max_images: int = 50) -> List[Dict]:
        """Télécharge les images et retourne les chemins locaux"""
        saved_images = []

        # Limiter le nombre d'images
        images_to_download = images[:max_images]

        self.logger.info(f"Téléchargement de {len(images_to_download)} images...")

        for i, img_info in enumerate(tqdm(images_to_download, desc="Images"), 1):
            try:
                img_url = img_info['url']
                img_alt = img_info['alt']

                # Déterminer l'extension
                if img_info['is_data_uri']:
                    # Data URI
                    ext = self._get_extension_from_data_uri(img_url)
                    img_path = os.path.join(self.output_dir, f"image_{i}.{ext}")
                    success = self._save_data_uri(img_url, img_path)
                else:
                    # URL HTTP
                    ext = self._get_extension_from_url(img_url)
                    img_path = os.path.join(self.output_dir, f"image_{i}.{ext}")
                    success = self._download_http_image(img_url, img_path)

                if success:
                    saved_images.append({
                        'path': img_path,
                        'alt': img_alt,
                        'original_url': img_url
                    })

            except Exception as e:
                self.logger.warning(f"Erreur lors du téléchargement de l'image {i}: {e}")
                continue

        self.logger.info(f"✓ {len(saved_images)} image(s) téléchargée(s)")
        return saved_images

    def _save_data_uri(self, data_uri: str, output_path: str) -> bool:
        """Sauvegarde une image depuis une Data URI"""
        try:
            # Format: data:image/png;base64,iVBOR...
            header, data = data_uri.split(',', 1)
            img_bytes = base64.b64decode(data)

            with open(output_path, 'wb') as f:
                f.write(img_bytes)

            return True
        except Exception as e:
            self.logger.warning(f"Erreur Data URI: {e}")
            return False

    def _download_http_image(self, url: str, output_path: str, timeout: int = 10) -> bool:
        """Télécharge une image depuis une URL HTTP"""
        try:
            response = self.session.get(url, stream=True, timeout=timeout)
            response.raise_for_status()

            with open(output_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)

            return True
        except Exception as e:
            self.logger.warning(f"Erreur téléchargement HTTP: {e}")
            return False

    def _get_extension_from_data_uri(self, data_uri: str) -> str:
        """Extrait l'extension depuis une Data URI"""
        try:
            mime_type = data_uri.split(';')[0].split(':')[1]
            return mime_type.split('/')[-1]
        except:
            return 'png'

    def _get_extension_from_url(self, url: str) -> str:
        """Extrait l'extension depuis une URL"""
        try:
            path = urlparse(url).path
            ext = os.path.splitext(path)[1][1:]
            if ext.lower() in ['jpg', 'jpeg', 'png', 'gif', 'webp', 'svg', 'bmp']:
                return ext.lower()
        except:
            pass
        return 'jpg'


# ==================== EXPORT MARKDOWN ====================

class MarkdownExporter:
    """Exportateur Markdown amélioré"""

    def __init__(self, logger: logging.Logger):
        self.logger = logger

    def export(self, content: Dict, output_file: str, include_images: bool = True):
        """Génère un fichier Markdown"""
        self.logger.info(f"Génération du fichier Markdown: {output_file}")

        with open(output_file, 'w', encoding='utf-8') as f:
            # Titre et métadonnées
            f.write(f"# {content['title']}\n\n")
            f.write(f"**URL source**: {content['url']}\n\n")
            f.write(f"**Date d'extraction**: {content['extraction_time']}\n\n")
            f.write("---\n\n")

            # Métadonnées
            if content.get('metadata'):
                f.write("## 📋 Métadonnées\n\n")
                for key, value in content['metadata'].items():
                    f.write(f"- **{key}**: {value}\n")
                f.write("\n---\n\n")

            # Texte principal
            if content.get('text'):
                f.write("## 📄 Contenu Textuel\n\n")
                f.write(f"{content['text']}\n\n")
                f.write("---\n\n")

            # Tableaux
            if content.get('tables'):
                f.write(f"## 📊 Tableaux ({len(content['tables'])})\n\n")
                for i, table in enumerate(content['tables'], 1):
                    f.write(f"### {table['title']}\n\n")
                    rows = table['rows']

                    if rows:
                        # En-tête
                        if len(rows) > 0:
                            f.write("| " + " | ".join(rows[0]) + " |\n")
                            f.write("|" + "|".join(["---"] * len(rows[0])) + "|\n")

                        # Données
                        for row in rows[1:]:
                            # S'assurer que toutes les lignes ont le même nombre de cellules
                            while len(row) < len(rows[0]):
                                row.append("")
                            f.write("| " + " | ".join(row[:len(rows[0])]) + " |\n")
                        f.write("\n")

                f.write("---\n\n")

            # Images
            if include_images and content.get('images'):
                f.write(f"## 🖼️ Images ({len(content['images'])})\n\n")
                for i, img in enumerate(content['images'], 1):
                    alt = img.get('alt', f"Image {i}")

                    if 'path' in img:
                        # Image téléchargée
                        rel_path = os.path.basename(img['path'])
                        f.write(f"### Image {i}: {alt}\n\n")
                        f.write(f"![{alt}]({self.config.IMAGES_DIR}/{rel_path})\n\n")
                    else:
                        # URL seulement
                        f.write(f"### Image {i}: {alt}\n\n")
                        f.write(f"- URL: {img['url']}\n\n")

                f.write("---\n\n")

            # Liens
            if content.get('links') and len(content['links']) <= 100:  # Limiter si trop de liens
                f.write(f"## 🔗 Liens Extraits ({len(content['links'])})\n\n")
                for link in content['links'][:100]:
                    f.write(f"- [{link['text']}]({link['url']})\n")
                f.write("\n---\n\n")

            f.write("**Fin du document**\n")

        self.logger.info(f"✓ Fichier Markdown généré: {output_file}")


# ==================== EXPORT WORD ====================

class WordExporter:
    """Exportateur Word (.docx)"""

    def __init__(self, logger: logging.Logger):
        self.logger = logger

    def export(self, content: Dict, output_file: str, include_images: bool = True):
        """Génère un fichier Word"""
        self.logger.info(f"Génération du fichier Word: {output_file}")

        doc = Document()

        # Style du document
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Titre
        title = doc.add_heading(content['title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Métadonnées de base
        doc.add_paragraph(f"URL source: {content['url']}")
        doc.add_paragraph(f"Date d'extraction: {content['extraction_time']}")
        doc.add_paragraph("_" * 50)

        # Métadonnées détaillées
        if content.get('metadata'):
            doc.add_heading('📋 Métadonnées', level=1)
            for key, value in content['metadata'].items():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
            doc.add_paragraph("_" * 50)

        # Texte principal
        if content.get('text'):
            doc.add_heading('📄 Contenu Textuel', level=1)

            # Diviser en paragraphes
            paragraphs = content['text'].split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

            doc.add_paragraph("_" * 50)

        # Tableaux
        if content.get('tables'):
            doc.add_heading(f'📊 Tableaux ({len(content["tables"])})', level=1)

            for i, table_data in enumerate(content['tables'], 1):
                doc.add_heading(table_data['title'], level=2)

                rows = table_data['rows']
                if rows and rows[0]:
                    # Créer le tableau Word
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'

                    # Remplir les cellules
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = str(cell_data)

                            # En-tête en gras
                            if row_idx == 0:
                                cell.paragraphs[0].runs[0].font.bold = True

                    doc.add_paragraph()  # Espace après le tableau

            doc.add_paragraph("_" * 50)

        # Images
        if include_images and content.get('images'):
            doc.add_heading(f'🖼️ Images ({len(content["images"])})', level=1)

            for i, img in enumerate(content['images'], 1):
                alt = img.get('alt', f"Image {i}")
                doc.add_heading(f"Image {i}: {alt}", level=2)

                # Ajouter l'image si elle existe localement
                if 'path' in img and os.path.exists(img['path']):
                    try:
                        # Limiter la largeur à 6 pouces
                        doc.add_picture(img['path'], width=Inches(6))
                        doc.add_paragraph()
                    except Exception as e:
                        self.logger.warning(f"Impossible d'ajouter l'image {i}: {e}")
                        doc.add_paragraph(f"URL: {img.get('original_url', 'N/A')}")
                else:
                    doc.add_paragraph(f"URL: {img.get('url', 'N/A')}")

            doc.add_paragraph("_" * 50)

        # Liens (limité)
        if content.get('links') and len(content['links']) <= 50:
            doc.add_heading(f'🔗 Liens Extraits ({len(content["links"])})', level=1)

            for link in content['links'][:50]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{link['text']}: ").bold = True
                p.add_run(link['url'])

        # Sauvegarder
        doc.save(output_file)
        self.logger.info(f"✓ Fichier Word généré: {output_file}")


# ==================== SCRAPER PRINCIPAL ====================

class AdvancedWebScraper:
    """Scraper web avancé - Classe principale"""

    def __init__(self, config: ScraperConfig = None):
        self.config = config or ScraperConfig()
        self.logger = setup_logging()

        self.web_driver = AdvancedWebDriver(self.config, self.logger)
        self.content_extractor = ContentExtractor(self.config, self.logger)
        self.markdown_exporter = MarkdownExporter(self.logger)
        self.word_exporter = WordExporter(self.logger)

        # Créer les dossiers de sortie
        os.makedirs(self.config.OUTPUT_DIR, exist_ok=True)
        images_dir = os.path.join(self.config.OUTPUT_DIR, self.config.IMAGES_DIR)
        os.makedirs(images_dir, exist_ok=True)

    def scrape(self, url: str, download_images: bool = True, export_formats: List[str] = None) -> Dict:
        """
        Scrappe une URL et exporte dans les formats spécifiés

        Args:
            url: URL à scraper
            download_images: Télécharger les images (True/False)
            export_formats: Liste de formats ['md', 'docx'] ou None pour les deux

        Returns:
            Dict avec le contenu extrait
        """
        export_formats = export_formats or ['md', 'docx']

        self.logger.info("=" * 70)
        self.logger.info(f"DÉBUT DU SCRAPING: {url}")
        self.logger.info("=" * 70)

        try:
            # 1. Créer le driver
            self.web_driver.create_driver()

            # 2. Charger la page
            if not self.web_driver.safe_get(url):
                raise Exception("Impossible de charger la page")

            # 3. Vérifier les blocages
            is_blocked, block_type = self.web_driver.detect_blocking()
            if is_blocked:
                self.logger.warning(f"⚠️ Blocage détecté: {block_type}")
                self.logger.warning("Tentative de contournement...")
                time.sleep(3)

            # 4. Gérer les pop-ups
            self.web_driver.handle_popups_and_cookies()
            time.sleep(1)

            # 5. Scroll pour lazy loading
            if self.config.SCROLL_TO_BOTTOM:
                self.web_driver.scroll_to_bottom()

            # 6. Extraire le contenu
            html = self.web_driver.driver.page_source
            content = self.content_extractor.extract_all(html, url)

            # 7. Télécharger les images
            if download_images and content.get('images'):
                images_dir = os.path.join(self.config.OUTPUT_DIR, self.config.IMAGES_DIR)
                downloader = ImageDownloader(images_dir, self.logger)
                saved_images = downloader.download_images(content['images'])
                content['images'] = saved_images

            # 8. Exporter
            base_name = self._generate_filename(url)

            if 'md' in export_formats:
                md_file = os.path.join(self.config.OUTPUT_DIR, f"{base_name}.md")
                self.markdown_exporter.export(content, md_file, download_images)

            if 'docx' in export_formats:
                docx_file = os.path.join(self.config.OUTPUT_DIR, f"{base_name}.docx")
                self.word_exporter.export(content, docx_file, download_images)

            self.logger.info("=" * 70)
            self.logger.info("✓ SCRAPING TERMINÉ AVEC SUCCÈS")
            self.logger.info("=" * 70)

            return content

        except Exception as e:
            self.logger.error(f"❌ ERREUR LORS DU SCRAPING: {e}", exc_info=True)
            raise

        finally:
            # Toujours fermer le driver
            self.web_driver.close()

    def _generate_filename(self, url: str) -> str:
        """Génère un nom de fichier basé sur l'URL"""
        parsed = urlparse(url)
        domain = parsed.netloc.replace('www.', '').replace('.', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f"{domain}_{timestamp}"


# ==================== FONCTION PRINCIPALE ====================

def main():
    """Point d'entrée principal du script"""

    print("\n" + "=" * 70)
    print("           WEB SCRAPER AVANCÉ - Version 2.0")
    print("=" * 70 + "\n")

    # Demander l'URL
    url = input("🌐 Entrez l'URL du site à scraper: ").strip()

    if not url:
        print("❌ URL invalide")
        return

    # Ajouter https:// si absent
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url

    # Options
    print("\n📋 OPTIONS:")
    download_images = input("Télécharger les images? (o/N): ").strip().lower() == 'o'

    export_md = input("Exporter en Markdown? (O/n): ").strip().lower() != 'n'
    export_docx = input("Exporter en Word? (O/n): ").strip().lower() != 'n'

    export_formats = []
    if export_md:
        export_formats.append('md')
    if export_docx:
        export_formats.append('docx')

    if not export_formats:
        print("❌ Aucun format d'export sélectionné")
        return

    # Configuration
    config = ScraperConfig()

    # Options avancées (optionnel)
    advanced = input("\nOptions avancées? (o/N): ").strip().lower() == 'o'
    if advanced:
        config.HEADLESS = input("Mode headless? (O/n): ").strip().lower() != 'n'
        config.SCROLL_TO_BOTTOM = input("Activer le scroll (lazy loading)? (O/n): ").strip().lower() != 'n'
        proxy = input("Proxy (laisser vide si aucun): ").strip()
        if proxy:
            # TODO: ajouter support proxy dans scrape()
            print("⚠️ Support proxy à implémenter dans create_driver()")

    # Lancer le scraping
    print("\n🚀 Démarrage du scraping...\n")

    try:
        scraper = AdvancedWebScraper(config)
        content = scraper.scrape(url, download_images, export_formats)

        print("\n✅ SCRAPING RÉUSSI!")
        print(f"\n📁 Résultats dans le dossier: {config.OUTPUT_DIR}/")
        print(f"   - Texte extrait: {len(content.get('text', ''))} caractères")
        print(f"   - Tableaux: {len(content.get('tables', []))}")
        print(f"   - Images: {len(content.get('images', []))}")
        print(f"   - Liens: {len(content.get('links', []))}")

    except KeyboardInterrupt:
        print("\n\n⚠️ Scraping interrompu par l'utilisateur")
    except Exception as e:
        print(f"\n❌ ERREUR: {e}")
        print("Consultez le fichier scraper.log pour plus de détails")


if __name__ == "__main__":
    main()
